from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path(r"C:\VuxNgoc\BaoCao_Travely_PhatTrienUngDungWeb.docx")
SCREENSHOT_DIR = Path(r"C:\VuxNgoc\travely_screenshots")


def set_run_font(run, size=13, bold=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), "Times New Roman")


def paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True, before=0, after=6):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.first_line_indent = Cm(1) if first_line else None


def add_p(doc, text="", bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True, before=0, after=6):
    p = doc.add_paragraph()
    paragraph_format(p, align=align, first_line=first_line, before=before, after=after)
    if text:
        run = p.add_run(text)
        set_run_font(run, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {min(level, 3)}"]
    paragraph_format(
        p,
        align=WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
        before=8 if level == 1 else 6,
        after=6,
    )
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=14 if level == 1 else 13, bold=True)
    if level == 1:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        paragraph_format(p, first_line=False, after=3)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        run = p.add_run("• " + item)
        set_run_font(run)


def add_numbered(doc, items):
    for index, item in enumerate(items, 1):
        p = doc.add_paragraph()
        paragraph_format(p, first_line=False, after=3)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        run = p.add_run(f"{index}. {item}")
        set_run_font(run)


def shade_cell(cell, fill="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    paragraph_format(p, align=align, first_line=False, after=0)
    run = p.add_run(str(text))
    set_run_font(run, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def should_center_cell(value):
    text = str(value).strip()
    return len(text) <= 18 or text.isdigit() or text.startswith(("GET ", "POST "))


def add_table(doc, caption, headers, rows):
    add_p(doc, caption, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, before=8, after=4)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 and should_center_cell(value) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], value, align=align)
    return table


def prepare_report_image(image_path, max_height=1120):
    if not image_path.exists():
        return None

    cropped_dir = image_path.parent / "report_crops"
    cropped_dir.mkdir(parents=True, exist_ok=True)
    output_path = cropped_dir / image_path.name

    with Image.open(image_path) as image:
        width, height = image.size
        crop_height = min(height, max_height)
        image.crop((0, 0, width, crop_height)).save(output_path)

    return output_path


def add_figure(doc, file_name, caption, width_cm=15.2):
    image_path = prepare_report_image(SCREENSHOT_DIR / file_name)
    if image_path is None:
        add_p(doc, f"Ảnh minh họa chưa được tạo: {file_name}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
        return

    p = doc.add_paragraph()
    paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, before=4, after=4)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    add_p(doc, caption, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, before=2, after=8)


def add_toc_field(doc):
    p = doc.add_paragraph()
    paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Mục lục sẽ tự động cập nhật khi mở bằng Microsoft Word."
    r.append(t)
    fld.append(r)
    run._r.append(fld)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    run._r.append(fld)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.5)
    add_page_number(section)

    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(14 if style_name == "Heading 1" else 13)
        style.font.bold = style_name.startswith("Heading")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    settings = doc.settings.element
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)
    for child in list(compat):
        if child.tag == qn("w:compatSetting") and child.get(qn("w:name")) == "compatibilityMode":
            compat.remove(child)
    compat_setting = OxmlElement("w:compatSetting")
    compat_setting.set(qn("w:name"), "compatibilityMode")
    compat_setting.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
    compat_setting.set(qn("w:val"), "15")
    compat.append(compat_setting)


def build_report():
    doc = Document()
    configure_document(doc)

    # Trang bìa
    for text in ("BỘ GIÁO DỤC VÀ ĐÀO TẠO", "TRƯỜNG ĐẠI HỌC THÀNH ĐÔNG", "KHOA CÔNG NGHỆ THÔNG TIN"):
        add_p(doc, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, after=2)
    for _ in range(3):
        doc.add_paragraph()
    add_p(doc, "BÁO CÁO MÔN PHÁT TRIỂN ỨNG DỤNG TRÊN NỀN WEB", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, after=10)
    add_p(doc, "ĐỀ TÀI:", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, after=0)
    p = add_p(doc, "XÂY DỰNG WEBSITE ĐẶT TOUR DU LỊCH TRAVELY", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, after=24)
    for run in p.runs:
        set_run_font(run, size=15, bold=True)

    info = [
        ("Giảng viên hướng dẫn", "ThS. Hoàng Anh Tuấn"),
        ("Sinh viên thực hiện", "Vũ Văn Ngọc"),
        ("Lớp", "1A11_CNTT14"),
        ("Ngành", "Công Nghệ Thông Tin"),
        ("Khóa", "14"),
    ]
    for label, value in info:
        p = add_p(doc, f"{label}: {value}", align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False, after=4)
        p.paragraph_format.left_indent = Cm(4.2)
    for _ in range(6):
        doc.add_paragraph()
    add_p(doc, "Hải Phòng, 2026", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    doc.add_page_break()

    # Lời cảm ơn, mục lục
    add_heading(doc, "LỜI CẢM ƠN", 1)
    thanks = [
        "Trong quá trình học tập môn Phát triển ứng dụng trên nền web và thực hiện đề tài xây dựng website đặt tour du lịch Travely, em đã nhận được sự hướng dẫn, hỗ trợ và động viên từ Nhà trường, Khoa Công nghệ Thông tin và các Thầy Cô giảng dạy.",
        "Em xin gửi lời cảm ơn chân thành tới ThS. Hoàng Anh Tuấn, người đã định hướng kiến thức, góp ý trong quá trình phân tích, thiết kế, xây dựng và kiểm thử sản phẩm. Những hướng dẫn của Thầy giúp em hiểu rõ hơn cách tổ chức một ứng dụng web theo mô hình MVC, cách xử lý dữ liệu và cách hoàn thiện giao diện theo yêu cầu thực tế.",
        "Mặc dù đã cố gắng hoàn thiện báo cáo và sản phẩm, bài làm khó tránh khỏi thiếu sót. Em mong nhận được ý kiến đóng góp của Quý Thầy Cô để có thể tiếp tục cải thiện trong các học phần sau.",
        "Em xin chân thành cảm ơn!",
    ]
    for text in thanks:
        add_p(doc, text)
    doc.add_page_break()

    add_heading(doc, "MỤC LỤC", 1)
    add_toc_field(doc)
    add_p(doc, "Ghi chú: Khi mở file bằng Microsoft Word, nhấn Ctrl + A rồi F9 để cập nhật số trang trong mục lục.", align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    doc.add_page_break()

    add_heading(doc, "DANH MỤC CỤM TỪ VIẾT TẮT", 1)
    add_table(doc, "Bảng 0.1. Danh mục cụm từ viết tắt", ["Từ viết tắt", "Ý nghĩa"], [
        ["MVC", "Model - View - Controller"],
        ["PHP", "Ngôn ngữ lập trình phía máy chủ"],
        ["PDO", "Thư viện truy cập cơ sở dữ liệu trong PHP"],
        ["MySQL", "Hệ quản trị cơ sở dữ liệu quan hệ"],
        ["XAMPP", "Bộ môi trường Apache, MySQL/MariaDB và PHP"],
        ["CSRF", "Cross-Site Request Forgery, hình thức giả mạo yêu cầu"],
    ])
    doc.add_page_break()

    # Mở đầu
    add_heading(doc, "MỞ ĐẦU", 1)
    add_heading(doc, "1. Lý do chọn đề tài", 2)
    add_p(doc, "Du lịch là lĩnh vực có nhu cầu tra cứu thông tin và đặt dịch vụ trực tuyến rất lớn. Người dùng thường muốn xem nhanh điểm đến, giá tour, lịch khởi hành, đánh giá, hình ảnh thực tế và gửi yêu cầu đặt tour mà không cần đến trực tiếp văn phòng. Vì vậy, một website đặt tour có giao diện rõ ràng, thao tác đơn giản và có khu vực quản trị là bài toán phù hợp với học phần Phát triển ứng dụng trên nền web.")
    add_p(doc, "Đề tài “Xây dựng website đặt tour du lịch Travely” được thực hiện nhằm vận dụng kiến thức về mô hình MVC, xử lý request/response, kết nối cơ sở dữ liệu, quản lý phiên đăng nhập, bảo vệ biểu mẫu bằng CSRF token và thiết kế giao diện web đáp ứng.")
    add_heading(doc, "2. Mục tiêu nghiên cứu", 2)
    add_bullets(doc, [
        "Xây dựng website đặt tour du lịch chạy trên XAMPP bằng PHP MVC và MySQL.",
        "Thiết kế giao diện công khai gồm trang chủ, danh sách tour, chi tiết tour, ưu đãi, liên hệ, đăng nhập và đăng ký.",
        "Xây dựng chức năng người dùng: đăng ký, đăng nhập, yêu thích tour, đặt tour và xem lịch sử booking.",
        "Xây dựng chức năng quản trị: dashboard, quản lý tour, booking, liên hệ, người dùng và phân quyền.",
        "Kiểm thử các chức năng chính, đánh giá ưu điểm, hạn chế và hướng phát triển sản phẩm.",
    ])
    add_heading(doc, "3. Đối tượng và phạm vi nghiên cứu", 2)
    add_p(doc, "Đối tượng nghiên cứu là ứng dụng web đặt tour du lịch Travely. Người dùng chính gồm khách truy cập, khách hàng đã đăng nhập và quản trị viên. Khách truy cập có thể xem tour, tìm kiếm, lọc thông tin và gửi liên hệ. Người dùng có tài khoản có thể đặt tour, lưu yêu thích và theo dõi booking. Quản trị viên có thể quản lý dữ liệu tour, booking, tin nhắn liên hệ và người dùng.")
    add_p(doc, "Phạm vi đề tài tập trung vào ứng dụng chạy cục bộ trên XAMPP, sử dụng PHP thuần theo mô hình MVC và MySQL. Các chức năng như thanh toán trực tuyến, gửi email tự động, xác thực OTP và triển khai lên máy chủ thật chưa nằm trong phạm vi bắt buộc của bài thực nghiệm.")
    add_heading(doc, "4. Phương pháp thực hiện", 2)
    add_p(doc, "Quá trình thực hiện kết hợp phương pháp phân tích yêu cầu, thiết kế cơ sở dữ liệu, tổ chức mã nguồn theo mô hình MVC, xây dựng giao diện bằng HTML/CSS/JavaScript, lập trình xử lý nghiệp vụ bằng PHP, kiểm thử chức năng bằng trình duyệt và script tự động.")
    add_heading(doc, "5. Kết cấu báo cáo", 2)
    add_p(doc, "Báo cáo gồm ba chương chính. Chương 1 trình bày cơ sở lý thuyết và công nghệ sử dụng. Chương 2 trình bày quá trình phân tích, thiết kế và xây dựng sản phẩm thực nghiệm. Chương 3 trình bày môi trường cài đặt, kiểm thử, demo và đánh giá kết quả.")
    doc.add_page_break()

    # Chương 1
    add_heading(doc, "CHƯƠNG 1. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ SỬ DỤNG", 1)
    add_heading(doc, "1.1. Tổng quan về phát triển ứng dụng web", 2)
    add_p(doc, "Ứng dụng web là phần mềm hoạt động thông qua trình duyệt và máy chủ. Người dùng gửi yêu cầu bằng HTTP, máy chủ xử lý logic nghiệp vụ, truy vấn cơ sở dữ liệu và trả kết quả dưới dạng HTML, CSS, JavaScript hoặc dữ liệu khác. Phần giao diện được xây dựng dựa trên các chuẩn HTML, CSS và JavaScript phổ biến trong tài liệu MDN Web Docs [4]. Ưu điểm của ứng dụng web là dễ truy cập, dễ cập nhật và phù hợp với nhiều thiết bị.")
    add_heading(doc, "1.2. Mô hình MVC", 2)
    add_p(doc, "MVC là mô hình kiến trúc chia ứng dụng thành ba thành phần: Model, View và Controller. Model chịu trách nhiệm biểu diễn dữ liệu, kiểm soát các thao tác truy vấn và cập nhật cơ sở dữ liệu. View chịu trách nhiệm hiển thị giao diện cho người dùng, nhận dữ liệu do Controller truyền sang và không nên chứa nhiều nghiệp vụ xử lý. Controller đóng vai trò trung gian, tiếp nhận request, kiểm tra điều kiện, gọi Model xử lý và chọn View phù hợp để trả kết quả.")
    add_p(doc, "Ưu điểm của mô hình MVC là tách biệt trách nhiệm giữa các lớp, giúp mã nguồn dễ đọc, dễ kiểm thử và dễ bảo trì. Khi cần sửa giao diện, lập trình viên chủ yếu làm việc với View; khi cần thay đổi truy vấn dữ liệu thì chỉnh Model; khi thay đổi luồng nghiệp vụ thì chỉnh Controller. Cách tổ chức này phù hợp với PHP vì PHP có thể xử lý request, session, cookie, form POST và kết nối cơ sở dữ liệu trực tiếp trên máy chủ [1].")
    add_table(doc, "Bảng 1.1. Vai trò ba thành phần trong mô hình MVC", ["Thành phần", "Vai trò", "Ví dụ trong Travely"], [
        ["Model", "Làm việc với dữ liệu, truy vấn MySQL, trả dữ liệu cho Controller", "Tour, Booking, User, Favorite, ContactMessage"],
        ["View", "Hiển thị HTML, form, bảng dữ liệu và ảnh giao diện", "home/index.php, tours/show.php, admin/bookings.php"],
        ["Controller", "Nhận request, kiểm tra quyền, gọi Model và chọn View", "TourController, BookingController, AdminController"],
    ])
    add_p(doc, "Trong website Travely, các Controller như HomeController, TourController, AuthController, BookingController, UserController, ContactController và AdminController phụ trách điều hướng nghiệp vụ. Các Model như Tour, Booking, User, Favorite và ContactMessage làm việc với cơ sở dữ liệu thông qua lớp Database sử dụng PDO [2].")
    add_heading(doc, "1.3. Mô hình máy khách - máy chủ", 2)
    add_p(doc, "Mô hình máy khách - máy chủ là nền tảng của hầu hết ứng dụng web. Máy khách là trình duyệt của người dùng, chịu trách nhiệm gửi yêu cầu và hiển thị giao diện. Máy chủ là Apache/PHP trong XAMPP, tiếp nhận yêu cầu, xử lý route, truy vấn MySQL và trả phản hồi. Trong phạm vi đề tài, XAMPP được sử dụng để mô phỏng môi trường chạy web đầy đủ trên máy cá nhân [5].")
    add_heading(doc, "1.4. Công nghệ và công cụ sử dụng", 2)
    add_table(doc, "Bảng 1.2. Nhóm công nghệ và công cụ sử dụng", ["Nhóm", "Công nghệ", "Vai trò trong đề tài"], [
        ["Ngôn ngữ lập trình", "PHP 8", "Xử lý request, session, controller, model và nghiệp vụ đặt tour"],
        ["Cơ sở dữ liệu", "MySQL/MariaDB", "Lưu người dùng, tour, booking, yêu thích và liên hệ [3]"],
        ["Kết nối dữ liệu", "PDO", "Thực thi truy vấn SQL an toàn bằng prepared statement [2]"],
        ["Giao diện", "HTML5, CSS3, JavaScript", "Xây dựng giao diện, hiệu ứng, preview ảnh URL và tương tác tab/gallery [4]"],
        ["Môi trường chạy", "XAMPP", "Cung cấp Apache, PHP và MySQL để chạy cục bộ [5]"],
        ["Kiểm thử", "PowerShell, curl, PHP lint", "Kiểm tra trang, đăng nhập, booking, admin và lỗi cú pháp"],
    ])
    add_heading(doc, "1.5. Bảo mật cơ bản trong ứng dụng web", 2)
    add_p(doc, "Website sử dụng session để lưu trạng thái đăng nhập, phân quyền admin bằng thuộc tính role trong bảng users và sử dụng CSRF token cho các form POST. Dữ liệu đầu ra được escape bằng hàm e() nhằm giảm nguy cơ XSS. Mật khẩu người dùng được lưu bằng password_hash và xác thực bằng password_verify theo cơ chế sẵn có của PHP [1].")
    add_heading(doc, "1.6. Nguyên tắc thiết kế giao diện", 2)
    add_p(doc, "Giao diện Travely được xây dựng theo định hướng du lịch hiện đại: hình ảnh lớn, màu sắc sáng, nút hành động rõ ràng và bố cục dễ đọc. Các phần không có chức năng thật được loại bỏ hoặc thay bằng liên kết hoạt động để tránh gây hiểu nhầm khi demo.")
    add_heading(doc, "1.7. Hệ quản trị cơ sở dữ liệu MySQL", 2)
    add_p(doc, "MySQL là hệ quản trị cơ sở dữ liệu quan hệ, thường được dùng trong các website PHP để lưu trữ dữ liệu có cấu trúc. Dữ liệu được tổ chức thành database, table, row và column. Mỗi bảng có các trường dữ liệu, khóa chính và có thể có quan hệ logic với bảng khác thông qua khóa ngoại hoặc trường liên kết. Trong đề tài Travely, MySQL lưu người dùng, tour, booking, danh sách yêu thích và tin nhắn liên hệ [3].")
    add_p(doc, "Các thao tác cơ bản với MySQL gồm tạo cơ sở dữ liệu, tạo bảng, thêm dữ liệu bằng INSERT, đọc dữ liệu bằng SELECT, cập nhật bằng UPDATE và xóa bằng DELETE. Khi triển khai cục bộ bằng XAMPP, sinh viên có thể import file database/travely.sql để tạo cấu trúc bảng và dữ liệu mẫu; khi cần sao lưu hoặc nộp bài, có thể export lại database ra file SQL.")
    add_table(doc, "Bảng 1.3. Một số thao tác MySQL dùng trong đề tài", ["Thao tác", "Mục đích", "Ví dụ áp dụng"], [
        ["CREATE", "Tạo database và bảng", "Tạo bảng users, tours, bookings"],
        ["SELECT", "Đọc dữ liệu", "Lấy danh sách tour, booking, liên hệ"],
        ["INSERT", "Thêm dữ liệu mới", "Tạo tài khoản, tạo booking, gửi liên hệ"],
        ["UPDATE", "Cập nhật dữ liệu", "Đổi trạng thái booking, sửa tour"],
        ["DELETE", "Xóa dữ liệu", "Xóa tour hoặc người dùng thường trong admin"],
        ["IMPORT/EXPORT", "Khởi tạo hoặc sao lưu dữ liệu", "Import file travely.sql khi demo"],
    ])
    add_heading(doc, "1.8. Lập trình MVC PHP kết nối MySQL", 2)
    add_p(doc, "Quy trình xây dựng một website PHP theo mô hình MVC thường bắt đầu từ điểm vào public/index.php. Tệp này nạp cấu hình, helper, autoload và route. Sau đó lớp điều phối route xác định Controller/action phù hợp. Controller kiểm tra request, gọi Model để làm việc với cơ sở dữ liệu, rồi truyền dữ liệu sang View để hiển thị.")
    add_p(doc, "Trong Travely, lớp Database tạo kết nối PDO tới MySQL. Các Model không nối chuỗi SQL trực tiếp với dữ liệu người dùng mà sử dụng prepared statement để giảm rủi ro SQL injection [2]. Cách tách lớp này giúp phần truy vấn dữ liệu nằm trong Model, phần xử lý luồng nằm trong Controller và phần hiển thị nằm trong View.")
    add_numbered(doc, [
        "Người dùng truy cập URL hoặc gửi form từ trình duyệt.",
        "public/index.php khởi tạo ứng dụng và nạp routes/web.php.",
        "Lớp App so khớp URL với route đã khai báo.",
        "Controller nhận request, kiểm tra quyền đăng nhập hoặc CSRF nếu cần.",
        "Controller gọi Model để SELECT, INSERT, UPDATE hoặc DELETE dữ liệu trong MySQL.",
        "Controller truyền dữ liệu sang View và View render HTML trả về trình duyệt.",
    ])
    doc.add_page_break()

    # Chương 2
    add_heading(doc, "CHƯƠNG 2. PHÂN TÍCH, THIẾT KẾ VÀ XÂY DỰNG SẢN PHẨM THỰC NGHIỆM", 1)
    add_heading(doc, "2.1. Giới thiệu bài toán", 2)
    add_p(doc, "Travely là website đặt tour du lịch cho phép người dùng tìm kiếm hành trình trong nước và nước ngoài, xem chi tiết tour, lưu tour yêu thích, đặt tour theo ngày khởi hành và gửi liên hệ tư vấn. Hệ thống có khu vực quản trị riêng để quản lý tour, booking, tin nhắn liên hệ và người dùng.")
    add_heading(doc, "2.2. Yêu cầu chức năng và phi chức năng", 2)
    add_table(doc, "Bảng 2.1. Yêu cầu chức năng của hệ thống", ["Nhóm người dùng", "Chức năng"], [
        ["Khách truy cập", "Xem trang chủ, xem danh sách tour, lọc tour, xem chi tiết, xem ưu đãi và gửi form liên hệ"],
        ["Người dùng đã đăng nhập", "Đặt tour, lưu yêu thích, xem thông tin tài khoản và lịch sử booking"],
        ["Quản trị viên", "Xem dashboard, thêm/sửa/xóa tour, quản lý booking, quản lý liên hệ, quản lý người dùng và phân quyền"],
    ])
    add_table(doc, "Bảng 2.2. Yêu cầu phi chức năng", ["Yêu cầu", "Mô tả"], [
        ["Dễ sử dụng", "Các thao tác chính được đặt ở vị trí rõ ràng, form có nhãn và kiểm tra dữ liệu"],
        ["Bảo mật cơ bản", "Sử dụng session, phân quyền, CSRF token, escape output và hash mật khẩu"],
        ["Tính ổn định", "Các route chính trả HTTP 200, không có lỗi cú pháp PHP trong mã nguồn"],
        ["Tính mở rộng", "Mã nguồn chia theo MVC, dễ bổ sung bảng dữ liệu và module mới"],
    ])
    add_heading(doc, "2.3. Luồng xử lý tổng thể", 2)
    add_p(doc, "Luồng xử lý cơ bản của website bắt đầu khi người dùng truy cập trình duyệt. Request được Apache chuyển tới public/index.php. Lớp App đọc route trong routes/web.php, xác định Controller và action tương ứng. Controller gọi Model để lấy hoặc ghi dữ liệu trong MySQL. Sau đó View được render trong layout và trả về trình duyệt.")
    add_table(doc, "Bảng 2.3. Luồng xử lý chính", ["Bước", "Mô tả"], [
        ["1", "Người dùng truy cập URL hoặc gửi form từ trình duyệt"],
        ["2", "public/index.php nạp cấu hình, helper, autoload và route"],
        ["3", "App dispatch request tới Controller phù hợp"],
        ["4", "Controller kiểm tra quyền, CSRF và gọi Model xử lý dữ liệu"],
        ["5", "Model truy vấn MySQL bằng PDO prepared statement"],
        ["6", "View hiển thị dữ liệu trong layout main hoặc admin"],
    ])
    add_heading(doc, "2.3.1. Các thành phần MVC trong website Travely", 3)
    add_p(doc, "Website Travely được tổ chức theo đúng hướng tách trách nhiệm của mô hình MVC. Nhóm Controller nhận và xử lý yêu cầu, nhóm Model làm việc với cơ sở dữ liệu, còn nhóm View hiển thị giao diện. Việc liệt kê rõ từng thành phần giúp quá trình thuyết trình và bảo vệ bài dễ theo dõi hơn.")
    add_table(doc, "Bảng 2.4. Danh sách Controller và chức năng", ["Controller", "Chức năng chính"], [
        ["HomeController", "Hiển thị trang chủ và các dữ liệu nổi bật"],
        ["TourController", "Hiển thị danh sách tour, lọc tour, tour trong nước, tour nước ngoài, ưu đãi và chi tiết tour"],
        ["AuthController", "Xử lý đăng nhập, đăng ký, đăng xuất và điều hướng sau xác thực"],
        ["BookingController", "Kiểm tra dữ liệu và tạo booking cho người dùng đã đăng nhập"],
        ["UserController", "Hiển thị trang tài khoản, lịch sử booking và xử lý yêu thích tour"],
        ["ContactController", "Hiển thị form liên hệ và lưu tin nhắn tư vấn vào cơ sở dữ liệu"],
        ["AdminController", "Quản lý dashboard, tour, booking, liên hệ, người dùng và phân quyền"],
    ])
    add_table(doc, "Bảng 2.5. Danh sách Model và chức năng", ["Model", "Chức năng chính"], [
        ["Tour", "Truy vấn danh sách tour, chi tiết tour, lọc tour, tạo/sửa/xóa tour trong admin"],
        ["Booking", "Tạo booking, lấy danh sách booking, tính doanh thu và cập nhật trạng thái"],
        ["User", "Tạo tài khoản, tìm người dùng, cập nhật quyền và xóa người dùng thường"],
        ["Favorite", "Lưu và xóa tour yêu thích của người dùng"],
        ["ContactMessage", "Lưu tin nhắn liên hệ, đếm tin chưa đọc và cập nhật trạng thái xử lý"],
        ["Database", "Tạo kết nối PDO tới MySQL để các Model sử dụng chung"],
    ])
    add_table(doc, "Bảng 2.6. Nhóm View và chức năng", ["Nhóm View", "Chức năng hiển thị"], [
        ["layouts/main.php", "Khung giao diện public gồm header, footer, thông báo và trợ lý AI"],
        ["layouts/admin.php", "Khung giao diện quản trị và menu điều hướng admin"],
        ["home, tours, contact", "Trang chủ, danh sách tour, chi tiết tour, ưu đãi và liên hệ"],
        ["auth", "Trang đăng nhập và đăng ký"],
        ["user", "Trang tài khoản người dùng và lịch sử booking"],
        ["admin", "Dashboard, quản lý tour, booking, liên hệ và người dùng"],
        ["partials", "Các thành phần dùng lại như tour card và trợ lý AI"],
    ])
    add_heading(doc, "2.4. Thiết kế cơ sở dữ liệu", 2)
    add_p(doc, "Cơ sở dữ liệu travely_cinematic_mvc gồm các bảng chính: users, tours, bookings, favorites và contact_messages. Bảng bookings liên kết users và tours, bảng favorites lưu tour yêu thích, bảng contact_messages lưu yêu cầu liên hệ.")
    add_table(doc, "Bảng 2.7. Các bảng dữ liệu chính", ["Bảng", "Vai trò", "Trường tiêu biểu"], [
        ["users", "Lưu tài khoản người dùng và quản trị", "id, name, email, phone, password, role"],
        ["tours", "Lưu thông tin tour du lịch", "title, slug, type, region, price, gallery, start_dates, status"],
        ["bookings", "Lưu yêu cầu đặt tour", "tour_id, user_id, start_date, guests, total_price, status"],
        ["favorites", "Lưu danh sách tour yêu thích", "user_id, tour_id"],
        ["contact_messages", "Lưu tin nhắn liên hệ", "name, phone, email, subject, message, status"],
    ])
    add_heading(doc, "2.5. Thiết kế giao diện và nhóm chức năng", 2)
    add_p(doc, "Giao diện công khai gồm trang chủ, danh sách tour, tour trong nước, tour nước ngoài, trang ưu đãi, trang chi tiết tour, liên hệ, đăng nhập và đăng ký. Trang chi tiết tour hiển thị gallery ảnh, thông tin tour, giá, lịch khởi hành, form đặt tour, điểm nổi bật và lịch trình.")
    add_p(doc, "Giao diện quản trị dùng layout riêng với thanh điều hướng, dashboard và các bảng quản lý. Form thêm/sửa tour cho phép dán URL ảnh online, có preview ảnh thumbnail và hero ngay trong giao diện admin.")
    add_table(doc, "Bảng 2.8. Các module trong khu vực quản trị", ["Module", "Chức năng"], [
        ["Dashboard", "Hiển thị số tour, booking, doanh thu, liên hệ mới và dữ liệu gần đây"],
        ["Quản lý tour", "Thêm, sửa, xóa tour; nhập URL ảnh online; quản lý trạng thái active/draft"],
        ["Booking", "Xem booking và cập nhật trạng thái pending, confirmed, completed, cancelled"],
        ["Liên hệ", "Xem tin nhắn từ form liên hệ và cập nhật trạng thái xử lý"],
        ["Người dùng", "Xem tài khoản, phân quyền admin/user và xóa người dùng thường"],
    ])
    add_heading(doc, "2.6. Phân tích chức năng đặt tour", 2)
    add_p(doc, "Chức năng đặt tour yêu cầu người dùng đăng nhập. Khi gửi form, hệ thống kiểm tra CSRF token, kiểm tra tour tồn tại, kiểm tra ngày khởi hành không nhỏ hơn ngày hiện tại và phải thuộc danh sách start_dates của tour. Số khách được giới hạn từ 1 đến 20.")
    add_table(doc, "Bảng 2.9. Ràng buộc dữ liệu khi đặt tour", ["Dữ liệu", "Ràng buộc"], [
        ["tour_id", "Phải tồn tại trong bảng tours"],
        ["start_date", "Định dạng YYYY-MM-DD, không trước ngày hiện tại và thuộc lịch mở bán của tour"],
        ["guests", "Từ 1 đến 20 người"],
        ["full_name, phone", "Không được để trống"],
        ["email", "Đúng định dạng email"],
        ["total_price", "Tính tự động từ giá tour và số khách"],
    ])
    add_heading(doc, "2.7. Phân tích chức năng quản lý booking", 2)
    add_p(doc, "Quản trị viên có thể đổi trạng thái booking từ chờ xác nhận sang đã xác nhận, hoàn tất hoặc đã hủy. Khi booking đã ở trạng thái completed hoặc cancelled thì hệ thống không cho thay đổi trạng thái nữa. Giao diện admin cũng khóa select và chỉ hiển thị nhãn trạng thái để tránh thao tác nhầm.")
    add_heading(doc, "2.8. Tổ chức mã nguồn", 2)
    add_table(doc, "Bảng 2.10. Vai trò các thư mục mã nguồn", ["Thư mục/Tệp", "Vai trò"], [
        ["public/index.php", "Điểm vào của ứng dụng, nạp cấu hình và route"],
        ["routes/web.php", "Khai báo các route GET/POST"],
        ["app/core", "Chứa App, Controller, Database, Auth và helper"],
        ["app/controllers", "Xử lý nghiệp vụ theo từng nhóm chức năng"],
        ["app/models", "Truy vấn và cập nhật dữ liệu MySQL"],
        ["app/views", "Giao diện public, user và admin"],
        ["public/assets", "CSS, JavaScript và tài nguyên giao diện"],
        ["database/travely.sql", "Script tạo database và dữ liệu mẫu"],
    ])
    add_heading(doc, "2.9. Một số quyết định kỹ thuật", 2)
    add_bullets(doc, [
        "Sử dụng PHP MVC tự viết thay vì framework lớn để phù hợp mục tiêu học tập và dễ giải thích mã nguồn.",
        "Sử dụng PDO prepared statement để hạn chế lỗi SQL injection.",
        "Sử dụng CSRF token cho các form POST quan trọng.",
        "Dùng ảnh online qua URL để admin dễ thêm/sửa tour mà không cần chức năng upload file.",
        "Dùng script kiểm thử tự động để đảm bảo các chức năng chính vẫn hoạt động sau mỗi lần chỉnh sửa.",
    ])
    doc.add_page_break()

    # Chương 3
    add_heading(doc, "CHƯƠNG 3. CÀI ĐẶT, KIỂM THỬ, DEMO VÀ ĐÁNH GIÁ", 1)
    add_heading(doc, "3.1. Môi trường cài đặt", 2)
    add_table(doc, "Bảng 3.1. Môi trường chạy thử sản phẩm", ["Thành phần", "Thông tin"], [
        ["Hệ điều hành", "Windows"],
        ["Máy chủ web", "Apache trong XAMPP"],
        ["Ngôn ngữ", "PHP 8.2"],
        ["Cơ sở dữ liệu", "MySQL/MariaDB trong XAMPP"],
        ["Trình duyệt kiểm thử", "Google Chrome"],
        ["Đường dẫn chạy", "http://localhost/travely-cinematic/public/"],
    ])
    add_heading(doc, "3.2. Quy trình triển khai cục bộ", 2)
    add_numbered(doc, [
        r"Đặt thư mục dự án vào C:\xampp\htdocs\travely-cinematic.",
        "Bật Apache và MySQL trong XAMPP Control Panel.",
        "Import file database/travely.sql vào MySQL để tạo cơ sở dữ liệu và dữ liệu mẫu.",
        "Mở trình duyệt và truy cập http://localhost/travely-cinematic/public/.",
        "Đăng nhập admin bằng tài khoản admin và mật khẩu 123456 để kiểm tra khu vực quản trị.",
    ])
    add_heading(doc, "3.3. Kịch bản demo sản phẩm", 2)
    add_table(doc, "Bảng 3.2. Kịch bản demo chính", ["STT", "Kịch bản", "Kết quả mong đợi"], [
        ["1", "Mở trang chủ", "Hiển thị hero, form tìm tour, tour nổi bật và ưu đãi"],
        ["2", "Tìm kiếm tour Sapa", "Danh sách trả về đúng tour Sapa"],
        ["3", "Xem chi tiết tour", "Hiển thị gallery, thông tin tour, lịch khởi hành và form đặt tour"],
        ["4", "Đăng ký/đăng nhập người dùng", "Tạo tài khoản và truy cập trang tài khoản"],
        ["5", "Đặt tour", "Tạo booking pending và hiển thị trong tài khoản"],
        ["6", "Đăng nhập admin", "Truy cập dashboard quản trị"],
        ["7", "Cập nhật trạng thái booking", "Trạng thái booking được lưu trong cơ sở dữ liệu"],
        ["8", "Thêm/sửa tour", "Tour được lưu, ảnh online hiển thị qua URL"],
    ])
    add_heading(doc, "3.4. Kiểm thử chức năng", 2)
    add_p(doc, "Sản phẩm được kiểm thử bằng script tự động sử dụng curl, MySQL CLI và các thao tác HTTP. Script kiểm tra trang public, đăng nhập người dùng, yêu thích tour, tạo booking, gửi liên hệ, đăng nhập admin, cập nhật trạng thái booking/liên hệ, thêm/sửa/xóa tour. Kết quả kiểm thử gần nhất đạt 25/25 trường hợp thành công.")
    add_table(doc, "Bảng 3.3. Kết quả kiểm thử tổng hợp", ["Nhóm kiểm thử", "Số ca", "Kết quả"], [
        ["Trang public", "9", "Đạt"],
        ["Tìm kiếm tour", "1", "Đạt"],
        ["Người dùng và tài khoản", "2", "Đạt"],
        ["Yêu thích và booking", "2", "Đạt"],
        ["Liên hệ", "1", "Đạt"],
        ["Trang admin", "6", "Đạt"],
        ["Cập nhật dữ liệu admin", "4", "Đạt"],
        ["Tổng cộng", "25", "25/25 đạt"],
    ])
    add_heading(doc, "3.5. Kiểm thử dữ liệu đầu vào", 2)
    add_table(doc, "Bảng 3.4. Kiểm thử ràng buộc dữ liệu", ["Chức năng", "Dữ liệu kiểm thử", "Kết quả"], [
        ["Đặt tour", "Ngày đi trước ngày hiện tại", "Không tạo booking, hiển thị thông báo lỗi"],
        ["Đặt tour", "Ngày không thuộc lịch tour", "Không tạo booking"],
        ["Đặt tour", "Số khách 99", "Không tạo booking do vượt giới hạn"],
        ["Admin booking", "Booking đã completed rồi đổi về confirmed", "Không đổi được trạng thái"],
        ["Admin tour", "Thêm tour trùng slug", "Không fatal, báo lỗi slug đã tồn tại"],
        ["Liên hệ", "Email sai định dạng", "Không lưu và yêu cầu nhập lại"],
    ])
    add_heading(doc, "3.6. Mô tả giao diện demo", 2)
    add_p(doc, "Trang chủ sử dụng ảnh hero du lịch, form tìm tour và các khối tour nổi bật. Trang danh sách tour cho phép lọc theo khu vực, chủ đề, khoảng giá, thời gian và sắp xếp. Trang chi tiết tour có gallery ảnh, thông tin giá, lịch khởi hành, yêu thích và form đặt tour.")
    add_p(doc, "Trang đăng nhập và đăng ký được thiết kế thống nhất theo phong cách cinematic. Form admin thêm/sửa tour cho phép dán URL ảnh online, có preview thumbnail và hero để quản trị viên kiểm tra nhanh trước khi lưu.")
    add_heading(doc, "3.7. Đánh giá kết quả thực hiện", 2)
    add_table(doc, "Bảng 3.5. Đánh giá sản phẩm", ["Nội dung", "Đánh giá"], [
        ["Mức độ hoàn thành chức năng", "Hoàn thành các chức năng chính của website đặt tour và quản trị"],
        ["Tổ chức mã nguồn", "Phân chia theo MVC, dễ theo dõi và giải thích"],
        ["Cơ sở dữ liệu", "Có quan hệ giữa user, tour, booking, favorite và contact"],
        ["Giao diện", "Đồng bộ, hiện đại, phù hợp chủ đề du lịch"],
        ["Kiểm thử", "Có kiểm thử tự động và kiểm thử logic nghiệp vụ"],
    ])
    add_heading(doc, "3.8. Hạn chế", 2)
    add_bullets(doc, [
        "Chưa tích hợp thanh toán trực tuyến và xác nhận email tự động.",
        "Chưa có chức năng upload ảnh lên máy chủ, hiện dùng URL ảnh online.",
        "Chưa có phân trang nâng cao cho danh sách tour, booking và liên hệ.",
        "Chưa triển khai lên hosting thật, sản phẩm đang chạy ở môi trường cục bộ XAMPP.",
    ])
    add_heading(doc, "3.9. Hướng phát triển", 2)
    add_bullets(doc, [
        "Bổ sung thanh toán online qua VNPay, Momo hoặc cổng thanh toán ngân hàng.",
        "Bổ sung gửi email xác nhận booking cho khách hàng.",
        "Thêm upload ảnh có kiểm tra định dạng và dung lượng.",
        "Thêm phân trang, tìm kiếm nâng cao và lọc booking theo trạng thái/ngày đi.",
        "Triển khai lên hosting thật và cấu hình HTTPS.",
    ])

    add_heading(doc, "3.10. Hình ảnh minh họa giao diện website", 2)
    screenshot_pages = [
        (
            "Trang chủ",
            "01_trang_chu.png",
            "Trang chủ Travely với khu vực hero, form tìm kiếm tour và các khối nội dung chính.",
            "Trang chủ là điểm chạm đầu tiên của người dùng. Khu vực đầu trang sử dụng ảnh du lịch lớn để tạo cảm giác trực quan, đồng thời đặt form tìm tour ở vị trí dễ nhìn nhằm hỗ trợ thao tác nhanh. Các tour nổi bật, ưu đãi và lời kêu gọi hành động được sắp xếp theo thứ tự ưu tiên để người dùng có thể bắt đầu hành trình đặt tour ngay từ trang đầu.",
        ),
        (
            "Danh sách tour",
            "02_danh_sach_tour.png",
            "Danh sách tour có bộ lọc điểm đến, ngày đi, người đi, số khách và khoảng giá.",
            "Trang danh sách tour cho phép người dùng duyệt toàn bộ sản phẩm đang mở bán. Các bộ lọc được giữ lại ở phần đầu trang để người dùng dễ thu hẹp kết quả theo nhu cầu. Mỗi tour hiển thị ảnh, tên, giá, thời gian và nút xem chi tiết, bảo đảm thông tin quan trọng được trình bày trước khi người dùng quyết định đặt tour.",
        ),
        (
            "Chi tiết tour",
            "03_chi_tiet_tour.png",
            "Trang chi tiết tour hiển thị gallery ảnh thật, mô tả, giá, lịch khởi hành và form đặt tour.",
            "Trang chi tiết tour là nơi tập trung thông tin quan trọng nhất trước khi đặt. Người dùng có thể xem ảnh điểm đến, đọc mô tả, kiểm tra lịch khởi hành, chọn số khách và gửi yêu cầu booking. Hệ thống chỉ cho phép chọn ngày hợp lệ nằm trong danh sách ngày mở bán của tour, từ đó hạn chế lỗi dữ liệu khi tạo booking.",
        ),
        (
            "Tour nước ngoài",
            "04_tour_nuoc_ngoai.png",
            "Trang tour nước ngoài chỉ giữ các khu vực Châu Á và Châu Âu theo phạm vi dữ liệu hiện có.",
            "Theo yêu cầu hoàn thiện sản phẩm, phần tour nước ngoài được tinh gọn theo dữ liệu thực tế đang có trong hệ thống. Website chỉ hiển thị các khu vực hoạt động được, tránh để những mục trang trí hoặc danh mục chưa có dữ liệu gây hiểu nhầm trong quá trình demo và chấm bài.",
        ),
        (
            "Liên hệ",
            "05_lien_he.png",
            "Trang liên hệ có form gửi yêu cầu tư vấn và thông tin liên lạc cơ bản.",
            "Trang liên hệ được giữ theo hướng gọn và có chức năng thật. Người dùng nhập họ tên, email, số điện thoại, tiêu đề và nội dung. Khi gửi thành công, dữ liệu được lưu vào bảng contact_messages để admin xử lý trong khu vực quản trị.",
        ),
        (
            "Đăng nhập",
            "06_dang_nhap.png",
            "Trang đăng nhập dùng chung phong cách hình ảnh với website và hỗ trợ tài khoản admin.",
            "Trang đăng nhập phục vụ cả người dùng thường và quản trị viên. Sau khi nhập thông tin hợp lệ, hệ thống kiểm tra mật khẩu bằng password_verify, lưu user_id vào session và điều hướng theo quyền. Nếu tài khoản là admin, người dùng được chuyển sang dashboard quản trị.",
        ),
        (
            "Đăng ký",
            "07_dang_ky.png",
            "Trang đăng ký được thiết kế đồng bộ với trang đăng nhập, gồm họ tên, email, số điện thoại và mật khẩu.",
            "Trang đăng ký cho phép khách tạo tài khoản để đặt tour và lưu tour yêu thích. Form có kiểm tra email, mật khẩu tối thiểu sáu ký tự và chống trùng email. Sau khi tạo tài khoản thành công, hệ thống tự đăng nhập và chuyển người dùng về trang tài khoản.",
        ),
        (
            "Dashboard admin",
            "08_admin_dashboard.png",
            "Dashboard admin hiển thị tổng quan tour, booking, doanh thu, liên hệ và dữ liệu gần đây.",
            "Khu vực admin được tách layout riêng so với trang public. Dashboard cung cấp thông tin tổng quan để quản trị viên nắm nhanh tình trạng vận hành: số tour đang bán, tổng booking, booking chờ xử lý, liên hệ mới, người dùng và doanh thu đã xác nhận hoặc hoàn tất.",
        ),
        (
            "Quản lý tour",
            "09_admin_tours.png",
            "Màn hình quản lý tour cho phép sửa, xóa và chuyển tới form thêm tour.",
            "Bảng quản lý tour hiển thị các tour đang có trong cơ sở dữ liệu. Các nút thao tác được căn chỉnh lại để không bị lệch giao diện. Form thêm/sửa tour sử dụng URL ảnh online, có preview để admin kiểm tra ảnh trước khi lưu, phù hợp yêu cầu demo nhanh mà không cần upload file.",
        ),
        (
            "Quản lý booking",
            "10_admin_bookings.png",
            "Màn hình quản lý booking hỗ trợ cập nhật trạng thái và khóa booking đã hoàn tất hoặc đã hủy.",
            "Admin có thể thay đổi trạng thái booking khi booking còn ở trạng thái chờ xác nhận hoặc đã xác nhận. Khi booking chuyển sang hoàn tất hoặc đã hủy, hệ thống khóa thao tác cập nhật để bảo toàn lịch sử xử lý, tránh trường hợp dữ liệu doanh thu hoặc trạng thái bị thay đổi không hợp lý.",
        ),
    ]

    for index, (title, file_name, caption, description) in enumerate(screenshot_pages, 1):
        if index > 1:
            doc.add_page_break()
        add_heading(doc, f"3.10.{index}. {title}", 3)
        add_p(doc, description)
        add_figure(doc, file_name, f"Hình 3.{index}. {caption}")

    add_heading(doc, "3.11. Thuyết minh chi tiết các luồng hoạt động", 2)
    detailed_pages = [
        (
            "Luồng người dùng truy cập website",
            [
                "Khi người dùng mở website, request đầu tiên đi vào public/index.php. Tệp này khởi tạo session, nạp cấu hình, nạp helper, đăng ký autoload và tạo đối tượng App để xử lý route. Cách tổ chức này giúp toàn bộ request có chung một điểm vào, thuận tiện cho việc kiểm soát luồng chạy.",
                "Sau khi route được xác định, controller tương ứng sẽ lấy dữ liệu từ model và truyền sang view. Layout main chịu trách nhiệm phần header, footer, liên kết CSS/JavaScript và các thành phần dùng chung. Nhờ đó mỗi view chỉ cần tập trung vào nội dung riêng của trang.",
            ],
        ),
        (
            "Luồng tìm kiếm và lọc tour",
            [
                "Chức năng tìm tour nhận dữ liệu từ form gồm từ khóa, khu vực, chủ đề, khoảng giá, số khách, ngày đi và cách sắp xếp. TourController gom các tham số hợp lệ rồi gọi Tour::all để truy vấn cơ sở dữ liệu. Cách làm này giữ phần xử lý SQL trong model, không để câu truy vấn phân tán trong view.",
                "Khi lọc theo ngày đi, hệ thống chỉ trả về các tour có start_dates chứa ngày được chọn. Điều này giúp kết quả tìm kiếm bám sát dữ liệu thật, tránh trường hợp người dùng chọn một ngày nhưng tour không mở bán vào ngày đó.",
            ],
        ),
        (
            "Luồng xem chi tiết tour",
            [
                "Trang chi tiết tour lấy dữ liệu theo slug để URL thân thiện và dễ đọc. Nếu slug không tồn tại, hệ thống trả về trang lỗi 404 thay vì hiển thị dữ liệu rỗng. Các thông tin như gallery, highlight, lịch trình và ngày khởi hành được giải mã từ dữ liệu lưu trong bảng tours.",
                "Form đặt tour nằm ngay trong trang chi tiết để người dùng không phải chuyển sang nhiều bước. Khi chưa đăng nhập, người dùng được điều hướng sang trang đăng nhập. Khi đã đăng nhập, form gửi dữ liệu booking đến BookingController để kiểm tra và lưu.",
            ],
        ),
        (
            "Luồng đăng ký tài khoản",
            [
                "Người dùng nhập họ tên, email, số điện thoại và mật khẩu. AuthController kiểm tra dữ liệu bắt buộc, kiểm tra định dạng email và độ dài mật khẩu. Nếu email đã tồn tại trong bảng users, hệ thống báo lỗi và không tạo tài khoản trùng.",
                "Khi dữ liệu hợp lệ, User::create hash mật khẩu bằng password_hash trước khi lưu. Sau đó Auth::loginById tạo session đăng nhập cho tài khoản mới, giúp người dùng có thể đặt tour ngay mà không phải đăng nhập lại.",
            ],
        ),
        (
            "Luồng đăng nhập và phân quyền",
            [
                "Chức năng đăng nhập tìm tài khoản theo định danh người dùng nhập, sau đó kiểm tra mật khẩu bằng password_verify. Với tài khoản admin, hệ thống hỗ trợ đăng nhập bằng tên ngắn admin để thuận tiện khi demo bài báo cáo.",
                "Sau khi đăng nhập, role trong bảng users quyết định nơi điều hướng. Người dùng thường vào trang tài khoản, còn quản trị viên vào dashboard admin. Các controller quản trị đều gọi requireAdmin để ngăn người dùng thường truy cập trái phép.",
            ],
        ),
        (
            "Luồng đặt tour",
            [
                "BookingController kiểm tra CSRF token, kiểm tra đăng nhập, xác định tour tồn tại và đọc thông tin giá. Sau đó hệ thống kiểm tra ngày đi không được nhỏ hơn ngày hiện tại, đồng thời phải thuộc danh sách ngày mở bán của tour. Số khách được giới hạn từ 1 đến 20.",
                "Tổng tiền được tính tự động bằng giá tour nhân số khách, không lấy trực tiếp từ dữ liệu client. Cách tính này giúp tránh việc người dùng sửa giá trong form. Booking sau khi tạo có trạng thái pending để admin xác nhận.",
            ],
        ),
        (
            "Luồng quản trị tour",
            [
                "Admin có thể thêm mới hoặc chỉnh sửa tour trong một form thống nhất. Trước khi lưu, hệ thống tạo slug từ tên tour hoặc dùng slug được nhập, sau đó kiểm tra trùng slug để tránh lỗi URL và lỗi ràng buộc duy nhất trong cơ sở dữ liệu.",
                "Các trường ảnh thumbnail và hero dùng URL online. Giao diện admin có phần preview ảnh để quản trị viên kiểm tra ngay chất lượng hiển thị trước khi lưu. Đây là giải pháp phù hợp với phạm vi bài tập vì đơn giản hơn upload file nhưng vẫn linh hoạt khi thay ảnh.",
            ],
        ),
        (
            "Luồng quản trị booking",
            [
                "Trang booking hiển thị danh sách đặt tour kèm khách hàng, ngày đi, tổng tiền và trạng thái. Admin có thể cập nhật trạng thái từ pending sang confirmed, completed hoặc cancelled tùy tình huống xử lý thực tế.",
                "Khi booking đã completed hoặc cancelled, cả giao diện và model đều không cho cập nhật tiếp. Việc kiểm tra ở cả hai lớp giúp tránh lỗi thao tác trên giao diện và cả lỗi khi có request POST gửi trực tiếp từ bên ngoài.",
            ],
        ),
        (
            "Luồng xử lý liên hệ",
            [
                "Form liên hệ cho phép khách gửi yêu cầu tư vấn mà không cần đăng nhập. ContactController kiểm tra họ tên, email, số điện thoại, tiêu đề và nội dung trước khi lưu vào bảng contact_messages. Admin có thể xem và cập nhật trạng thái xử lý trong khu vực quản trị.",
                "Phần nhận deal ở footer trước đây chỉ là form trang trí đã được đổi thành liên kết tới trang liên hệ. Thay đổi này giúp website không còn nút giả, đúng với yêu cầu chỉ giữ những chức năng thật sự hoạt động được.",
            ],
        ),
        (
            "Kiểm soát bảo mật cơ bản",
            [
                "Các form POST quan trọng đều có CSRF token để hạn chế request giả mạo. Dữ liệu hiển thị ra giao diện được escape bằng helper e nhằm giảm nguy cơ XSS. Mật khẩu không lưu dạng rõ mà được hash trước khi ghi vào cơ sở dữ liệu.",
                "Phân quyền được kiểm tra ở controller thay vì chỉ ẩn nút trên giao diện. Điều này quan trọng vì người dùng có thể tự nhập URL admin trên thanh địa chỉ. Nếu không có quyền admin, hệ thống trả lỗi hoặc chuyển hướng phù hợp.",
            ],
        ),
        (
            "Kiểm thử tự động",
            [
                "Script kiểm thử tự động được dùng để kiểm tra luồng chính sau mỗi lần chỉnh sửa. Các ca kiểm thử bao gồm truy cập trang public, đăng nhập, đăng ký, yêu thích tour, đặt tour, gửi liên hệ, đăng nhập admin và thao tác quản trị.",
                "Việc chạy kiểm thử giúp phát hiện nhanh lỗi hồi quy. Trong lần kiểm thử gần nhất, hệ thống đạt 25/25 ca, cho thấy các chức năng cốt lõi hoạt động ổn định ở môi trường local.",
            ],
        ),
        (
            "Đánh giá mức độ hoàn thiện",
            [
                "Website Travely đã đáp ứng mục tiêu của học phần Phát triển ứng dụng trên nền web: có giao diện public, có xử lý form, có cơ sở dữ liệu, có đăng nhập phân quyền, có quản trị và có kiểm thử chức năng. Sản phẩm không chỉ là giao diện tĩnh mà có luồng dữ liệu từ người dùng đến MySQL.",
                "Một số phần có thể phát triển thêm như thanh toán online, gửi email xác nhận và upload ảnh. Tuy nhiên trong phạm vi bài báo cáo qua môn, các chức năng chính đã đủ để trình bày quy trình phân tích, thiết kế, cài đặt và kiểm thử một ứng dụng web hoàn chỉnh.",
            ],
        ),
    ]

    detail_notes = {
        "Luồng người dùng truy cập website": [
            "Điểm cần chú ý của luồng này là mọi request đều đi qua cùng một cơ chế điều phối. Khi cần thêm trang mới, lập trình viên chỉ cần khai báo route trong routes/web.php, tạo controller/action và view tương ứng. Cách làm này phù hợp với phạm vi học phần vì giúp sinh viên nhìn rõ quan hệ giữa URL, controller, model và view.",
            "Trong quá trình demo, luồng truy cập được kiểm chứng bằng các trang public như trang chủ, danh sách tour, chi tiết tour, ưu đãi, liên hệ, đăng nhập và đăng ký. Các trang này đều trả về HTTP 200 khi môi trường XAMPP/MySQL hoạt động đúng.",
        ],
        "Luồng tìm kiếm và lọc tour": [
            "Bộ lọc tour được thiết kế theo hướng chỉ giữ các lựa chọn có dữ liệu thật. Ví dụ phần nước ngoài chỉ còn Châu Á và Châu Âu, tránh việc người dùng chọn Châu Mỹ hoặc Châu Phi nhưng hệ thống không có tour để hiển thị.",
            "Với bài báo cáo, đây là điểm thể hiện rõ yêu cầu 'cái nào hoạt động được mới cho vào'. Các lựa chọn trên giao diện cần đi kèm dữ liệu và xử lý ở backend, không nên chỉ xuất hiện như thành phần trang trí.",
        ],
        "Luồng xem chi tiết tour": [
            "Thông tin chi tiết tour được trình bày theo thứ tự ưu tiên: ảnh, tiêu đề, đánh giá, mô tả, giá, lịch khởi hành và form đặt tour. Cách sắp xếp này giúp người dùng có đủ dữ liệu trước khi đưa ra quyết định.",
            "Gallery sử dụng ảnh thật theo điểm đến để tăng độ tin cậy. Đây cũng là phần được chỉnh sửa nhiều trong quá trình hoàn thiện sản phẩm, vì ảnh sai điểm đến sẽ làm giảm chất lượng demo và tính thuyết phục của báo cáo.",
        ],
        "Luồng đăng ký tài khoản": [
            "Tài khoản sau khi đăng ký có role mặc định là user. Người dùng thường không được phép vào khu vực admin, nhưng vẫn có thể đặt tour, lưu tour yêu thích và xem thông tin booking của mình.",
            "Chức năng đăng ký được giữ lại vì đã hoạt động thật. Form không chỉ hiển thị giao diện mà còn tạo bản ghi trong bảng users, kiểm tra trùng email và tự đăng nhập sau khi tạo tài khoản.",
        ],
        "Luồng đăng nhập và phân quyền": [
            "Tài khoản admin được đổi về dạng dễ nhớ là admin với mật khẩu 123456 để thuận tiện khi giáo viên kiểm tra. Mặc dù vậy, trong mã nguồn vẫn dùng password_hash/password_verify để bảo đảm cách lưu mật khẩu đúng kỹ thuật.",
            "Phân quyền không phụ thuộc vào việc ẩn nút trên giao diện. AdminController yêu cầu quyền quản trị ngay khi khởi tạo, vì vậy người dùng thường không thể truy cập route admin bằng cách nhập trực tiếp URL.",
        ],
        "Luồng đặt tour": [
            "Ràng buộc ngày đi được xử lý ở phía máy chủ thay vì chỉ dựa vào thuộc tính min của input date. Điều này quan trọng vì dữ liệu gửi từ trình duyệt có thể bị sửa thủ công, nên máy chủ luôn phải kiểm tra lại.",
            "Sau khi tạo booking, trạng thái ban đầu là pending. Admin xác nhận hoặc hủy tùy tình trạng thực tế. Khi booking hoàn tất hoặc đã hủy, hệ thống khóa cập nhật để bảo toàn lịch sử xử lý.",
        ],
        "Luồng quản trị tour": [
            "Form quản trị tour dùng URL ảnh online thay cho upload file để giảm độ phức tạp nhưng vẫn đáp ứng nhu cầu thay ảnh nhanh. Phần preview giúp admin phát hiện URL sai hoặc ảnh không phù hợp trước khi lưu.",
            "Khi thêm tour mới, slug được kiểm tra trùng để tránh lỗi dữ liệu. Nếu trùng slug, hệ thống báo lỗi thân thiện thay vì để website phát sinh lỗi SQL hoặc lỗi trang trắng.",
        ],
        "Luồng quản trị booking": [
            "Màn hình booking là phần quan trọng khi chấm bài vì chứng minh dữ liệu người dùng gửi từ public có thể được admin xử lý ở backend. Đây là điểm khác biệt giữa website tĩnh và ứng dụng web có nghiệp vụ.",
            "Các trạng thái booking được dùng nhất quán giữa model, controller và view. Sự nhất quán này giúp báo cáo dễ giải thích, đồng thời giảm lỗi khi kiểm thử tự động.",
        ],
        "Luồng xử lý liên hệ": [
            "Form liên hệ giúp khách gửi yêu cầu tư vấn khi chưa muốn đăng ký tài khoản. Đây là chức năng phù hợp với website du lịch vì nhiều khách hàng cần hỏi trước về lịch trình, giá hoặc số lượng khách.",
            "Admin có thể cập nhật trạng thái tin nhắn để theo dõi đã xử lý hay chưa. Trong báo cáo, chức năng này được dùng để minh họa một luồng dữ liệu độc lập với booking.",
        ],
        "Kiểm soát bảo mật cơ bản": [
            "Các biện pháp bảo mật được chọn ở mức phù hợp với bài học: session, CSRF, hash mật khẩu, escape output và phân quyền theo role. Đây là các nội dung có thể giải thích rõ trong phần vấn đáp.",
            "Hệ thống chưa triển khai các cơ chế nâng cao như rate limit, captcha hoặc xác thực email, nhưng phần lõi đã đủ để tránh các lỗi phổ biến khi xử lý form và tài khoản.",
        ],
        "Kiểm thử tự động": [
            "Việc có script kiểm thử là điểm cộng cho bài báo cáo vì chứng minh chức năng không chỉ được kiểm tra bằng mắt. Script có thể chạy lại sau mỗi lần sửa để bảo đảm các luồng chính vẫn hoạt động.",
            "Các ca kiểm thử được chọn theo hướng thực tế: truy cập trang, đăng nhập, tạo booking, gửi liên hệ, vào admin và cập nhật dữ liệu. Đây là các luồng giáo viên thường kiểm tra khi chấm sản phẩm.",
        ],
        "Đánh giá mức độ hoàn thiện": [
            "So với yêu cầu của một bài qua môn, Travely đã có đầy đủ các thành phần chính của ứng dụng web: giao diện, route, controller, model, database, session, form POST, phân quyền và kiểm thử.",
            "Điểm cần lưu ý khi trình bày là nhấn mạnh sản phẩm đã loại bỏ các phần trang trí không hoạt động, ví dụ form nhận deal giả, các danh mục nước ngoài không có dữ liệu và các icon/khối thông tin gây nhiễu.",
        ],
    }

    detail_tables = {
        "Luồng người dùng truy cập website": [
            ["URL", "Controller", "Kết quả"],
            [["/", "HomeController@index", "Trang chủ"], ["/tours", "TourController@index", "Danh sách tour"], ["/contact", "ContactController@index", "Form liên hệ"]],
        ],
        "Luồng tìm kiếm và lọc tour": [
            ["Tham số", "Ý nghĩa", "Cách xử lý"],
            [["q", "Từ khóa", "Tìm theo tên/mô tả tour"], ["region", "Khu vực", "Lọc theo miền hoặc châu lục"], ["start_date", "Ngày đi", "So khớp với lịch mở bán"]],
        ],
        "Luồng xem chi tiết tour": [
            ["Thành phần", "Dữ liệu hiển thị", "Nguồn"],
            [["Gallery", "Ảnh hero và ảnh phụ", "Bảng tours"], ["Form booking", "Ngày đi và số khách", "start_dates, price"], ["Lịch trình", "Các ngày trong tour", "itinerary"]],
        ],
        "Luồng đăng ký tài khoản": [
            ["Trường", "Ràng buộc", "Kết quả hợp lệ"],
            [["Họ tên", "Không rỗng", "Lưu vào users.name"], ["Email", "Đúng định dạng, không trùng", "Lưu vào users.email"], ["Mật khẩu", "Tối thiểu 6 ký tự", "Hash trước khi lưu"]],
        ],
        "Luồng đăng nhập và phân quyền": [
            ["Role", "Trang sau đăng nhập", "Quyền"],
            [["user", "/account", "Đặt tour, yêu thích, xem booking"], ["admin", "/admin", "Quản lý tour, booking, liên hệ, người dùng"]],
        ],
        "Luồng đặt tour": [
            ["Kiểm tra", "Lý do", "Kết quả khi sai"],
            [["CSRF", "Chống request giả mạo", "Từ chối POST"], ["Ngày đi", "Không đặt quá khứ/sai lịch", "Không tạo booking"], ["Số khách", "Giới hạn nghiệp vụ", "Báo lỗi nhập liệu"]],
        ],
        "Luồng quản trị tour": [
            ["Dữ liệu", "Cách nhập", "Kiểm soát"],
            [["Slug", "Tự sinh hoặc nhập tay", "Không được trùng"], ["Ảnh", "URL online", "Preview trước khi lưu"], ["Trạng thái", "active/draft", "Quyết định hiển thị public"]],
        ],
        "Luồng quản trị booking": [
            ["Trạng thái", "Ý nghĩa", "Có cho đổi tiếp không"],
            [["pending", "Chờ xác nhận", "Có"], ["confirmed", "Đã xác nhận", "Có"], ["completed/cancelled", "Kết thúc xử lý", "Không"]],
        ],
        "Luồng xử lý liên hệ": [
            ["Bước", "Người thực hiện", "Kết quả"],
            [["Gửi form", "Khách truy cập", "Tạo contact_messages"], ["Xem tin", "Admin", "Đọc nội dung yêu cầu"], ["Cập nhật", "Admin", "Đánh dấu trạng thái xử lý"]],
        ],
        "Kiểm soát bảo mật cơ bản": [
            ["Cơ chế", "Vị trí", "Mục đích"],
            [["CSRF token", "Form POST", "Chống giả mạo request"], ["password_hash", "User model", "Không lưu mật khẩu rõ"], ["requireAdmin", "Controller admin", "Chặn truy cập trái quyền"]],
        ],
        "Kiểm thử tự động": [
            ["Nhóm", "Cách kiểm tra", "Kết quả"],
            [["Public", "HTTP request", "Trang trả 200"], ["Auth/booking", "Session + POST", "Tạo dữ liệu đúng"], ["Admin", "Đăng nhập admin + POST", "Cập nhật dữ liệu đúng"]],
        ],
        "Đánh giá mức độ hoàn thiện": [
            ["Tiêu chí", "Mức đạt", "Minh chứng"],
            [["Chức năng", "Đạt", "Có public, user, admin"], ["Dữ liệu", "Đạt", "Có MySQL và quan hệ bảng"], ["Demo", "Đạt", "Có ảnh chụp và test 25/25"]],
        ],
    }

    for index, (title, paragraphs) in enumerate(detailed_pages, 1):
        add_heading(doc, f"3.11.{index}. {title}", 3)
        for paragraph in paragraphs:
            add_p(doc, paragraph)
        for paragraph in detail_notes.get(title, []):
            add_p(doc, paragraph)
        if title in detail_tables:
            headers, rows = detail_tables[title]
            add_table(doc, f"Bảng 3.6.{index}. Tóm tắt {title.lower()}", headers, rows)

    add_heading(doc, "KẾT LUẬN", 1)
    add_p(doc, "Đề tài “Xây dựng website đặt tour du lịch Travely” đã vận dụng các kiến thức của môn Phát triển ứng dụng trên nền web vào một sản phẩm có dữ liệu và chức năng thực tế. Website được xây dựng theo mô hình PHP MVC, sử dụng MySQL để lưu dữ liệu, có giao diện public, tài khoản người dùng và khu vực quản trị.")
    add_p(doc, "Sản phẩm đã hoàn thành các chức năng phía người dùng gồm: xem trang chủ, xem danh sách tour, lọc tour, xem chi tiết tour, đăng ký, đăng nhập, lưu tour yêu thích, đặt tour theo lịch khởi hành hợp lệ, xem tài khoản và gửi liên hệ tư vấn. Các chức năng này chứng minh website không chỉ là giao diện tĩnh mà có xử lý dữ liệu và lưu trữ trong MySQL.")
    add_p(doc, "Sản phẩm cũng đã hoàn thành các chức năng phía quản trị gồm: xem dashboard tổng quan, thêm/sửa/xóa tour, sử dụng ảnh online cho thumbnail và hero, quản lý booking, khóa booking đã hoàn tất hoặc đã hủy, quản lý tin nhắn liên hệ, quản lý người dùng và phân quyền. Các kiểm thử chính đạt kết quả 25/25, đồng thời một số lỗi logic như đặt ngày không hợp lệ, đổi trạng thái booking đã hoàn tất và thêm tour trùng slug đã được xử lý.")
    add_p(doc, "Thông qua quá trình thực hiện, em hiểu rõ hơn cách tổ chức ứng dụng MVC, thiết kế bảng dữ liệu, xử lý form, bảo mật cơ bản bằng session/CSRF, kiểm thử chức năng và hoàn thiện giao diện theo yêu cầu người dùng.")

    add_heading(doc, "KIẾN NGHỊ", 1)
    add_p(doc, "Để website Travely có thể phát triển theo hướng ứng dụng thực tế hơn, em kiến nghị bổ sung một số chức năng ở các phiên bản tiếp theo. Trước hết, hệ thống nên tích hợp thanh toán trực tuyến qua VNPay, Momo hoặc cổng thanh toán ngân hàng để người dùng có thể đặt cọc tour ngay trên website. Tiếp theo, website nên có chức năng gửi email xác nhận booking, thông báo thay đổi trạng thái và nhắc lịch khởi hành.")
    add_p(doc, "Ngoài ra, hệ thống có thể bổ sung upload ảnh lên máy chủ, phân trang nâng cao, tìm kiếm booking theo trạng thái/ngày đi, thống kê doanh thu theo tháng và triển khai lên hosting thật có HTTPS. Những hướng phát triển này giúp website tối ưu hơn về trải nghiệm người dùng, quản trị dữ liệu và khả năng vận hành trong môi trường thực tế.")

    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    add_p(doc, "[1] PHP Manual, sử dụng trong các mục 1.4, 1.5 và phần xử lý PHP: https://www.php.net/manual/en/", first_line=False)
    add_p(doc, "[2] PHP PDO Manual, sử dụng trong các mục 1.2, 1.8 và phần kết nối MySQL: https://www.php.net/manual/en/book.pdo.php", first_line=False)
    add_p(doc, "[3] MySQL Documentation, sử dụng trong mục 1.7 và phần thiết kế cơ sở dữ liệu: https://dev.mysql.com/doc/", first_line=False)
    add_p(doc, "[4] MDN Web Docs về HTML, CSS và JavaScript, sử dụng trong mục 1.1, 1.4 và phần thiết kế giao diện: https://developer.mozilla.org/", first_line=False)
    add_p(doc, "[5] XAMPP Documentation, sử dụng trong mục 1.3, 1.4 và phần triển khai cục bộ: https://www.apachefriends.org/", first_line=False)

    add_heading(doc, "PHỤ LỤC", 1)
    add_heading(doc, "Phụ lục A. Một số route chính", 2)
    add_table(doc, "Bảng A.1. Route tiêu biểu trong website Travely", ["Route", "Chức năng"], [
        ["GET /", "Trang chủ"],
        ["GET /tours", "Danh sách tour"],
        ["GET /tour/{slug}", "Chi tiết tour"],
        ["POST /booking/store", "Tạo booking"],
        ["POST /favorite/toggle", "Thêm/xóa yêu thích"],
        ["GET /admin", "Dashboard admin"],
        ["POST /admin/tours/save", "Thêm hoặc sửa tour"],
        ["POST /admin/bookings/status/{id}", "Cập nhật trạng thái booking"],
    ])

    add_heading(doc, "Phụ lục B. Tài khoản và dữ liệu demo", 2)
    add_p(doc, "Phần phụ lục này ghi lại các thông tin dùng khi giáo viên chạy thử sản phẩm trên môi trường local. Các tài khoản demo chỉ dùng cho mục đích kiểm tra bài báo cáo, không dùng cho hệ thống triển khai thật.")
    add_table(doc, "Bảng B.1. Tài khoản demo", ["Loại tài khoản", "Tên đăng nhập", "Mật khẩu", "Mục đích"], [
        ["Quản trị viên", "admin", "123456", "Truy cập dashboard và các trang quản lý"],
        ["Người dùng thường", "Tạo tại trang đăng ký", "Tối thiểu 6 ký tự", "Đặt tour, lưu yêu thích và xem lịch sử booking"],
    ])
    add_table(doc, "Bảng B.2. Dữ liệu demo chính", ["Nhóm dữ liệu", "Ví dụ", "Ghi chú"], [
        ["Tour trong nước", "Hạ Long, Đà Nẵng - Hội An, Sapa, Phú Quốc", "Có ảnh đúng điểm đến và lịch khởi hành"],
        ["Tour nước ngoài", "Seoul, Tokyo, Zurich, Singapore, Bali", "Chỉ giữ Châu Á và Châu Âu theo phạm vi dữ liệu"],
        ["Booking", "Chờ xác nhận, đã xác nhận, hoàn tất, đã hủy", "Trạng thái hoàn tất/đã hủy không cho đổi lại"],
        ["Liên hệ", "Tin nhắn từ form liên hệ", "Admin có thể xem và cập nhật trạng thái"],
    ])

    add_heading(doc, "Phụ lục C. Ghi chú kiểm thử", 2)
    add_p(doc, "Khi kiểm thử lại website, cần bật Apache/MySQL trong XAMPP hoặc chạy máy chủ PHP tích hợp với đúng thư mục public. Sau đó truy cập trang chủ, đăng nhập admin, thử tạo booking và kiểm tra các trang quản trị. Nếu dữ liệu bị thay đổi nhiều lần trong quá trình demo, có thể import lại file database/travely.sql để đưa hệ thống về trạng thái ban đầu.")
    add_table(doc, "Bảng C.1. Nhóm kiểm thử nên chạy lại trước khi nộp", ["Nhóm", "Thao tác", "Kết quả cần đạt"], [
        ["Public", "Mở trang chủ, danh sách tour, chi tiết tour, liên hệ", "Trang hiển thị đúng và không lỗi 404/500"],
        ["Tài khoản", "Đăng ký, đăng nhập, đăng xuất", "Session hoạt động đúng"],
        ["Booking", "Chọn ngày hợp lệ và đặt tour", "Booking được tạo ở trạng thái chờ xác nhận"],
        ["Admin", "Đăng nhập admin và cập nhật booking", "Trạng thái được lưu, booking hoàn tất/đã hủy bị khóa"],
        ["Tour", "Thêm hoặc sửa tour bằng URL ảnh online", "Ảnh preview và hiển thị đúng ngoài public"],
    ])

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, size=14 if paragraph.style.name == "Heading 1" else 13, bold=paragraph.style.name.startswith("Heading") or run.bold)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_report()
    print(OUTPUT)
